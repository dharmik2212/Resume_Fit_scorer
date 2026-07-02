"""Extract readable text and candidate identity from uploaded resumes."""
import io
import logging
import re

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as exc:
        logger.error("PDF parsing error: %s", exc)
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        document = Document(io.BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as exc:
        logger.error("DOCX parsing error: %s", exc)
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif suffix == "docx":
        text = extract_text_from_docx(file_bytes)
    elif suffix == "txt":
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        return ""
    return clean_text(text)


def clean_text(text: str) -> str:
    """Normalize text while preserving lines used for candidate-name detection."""
    lines = []
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = re.sub(r"[^\x20-\x7E]", " ", line)
        line = re.sub(r"[ \t]+", " ", line).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


SECTION_HEADERS = [
    "experience", "work experience", "professional experience", "employment",
    "education", "academic", "qualifications",
    "skills", "technical skills", "core competencies", "expertise",
    "projects", "project experience",
    "summary", "professional summary", "profile", "objective",
    "certifications", "certificates",
    "publications", "research",
    "achievements", "awards", "honors",
    "languages", "interests", "volunteer",
]


def parse_resume_sections(text: str) -> dict[str, str]:
    """Split resume text into sections based on common headers."""
    lines = text.splitlines()
    sections: dict[str, list[str]] = {"header": []}
    current = "header"
    header_pattern = re.compile(
        r"^(?:\d{1,2}[./])?\s*(" + "|".join(re.escape(h) for h in SECTION_HEADERS) + r")s?\s*$",
        re.I,
    )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Check for short all-caps section headers
        lower = stripped.lower().rstrip(":")
        if lower in SECTION_HEADERS:
            current = lower
            sections.setdefault(current, [])
        elif header_pattern.match(stripped):
            h = stripped.lower().strip()
            current = h
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(stripped)

    # Trim to first section occurrence only
    result: dict[str, str] = {}
    seen = set()
    for key, val in sections.items():
        base = key.rstrip("s").rstrip(":")
        if base not in seen:
            result[base] = "\n".join(val)
            seen.add(base)
        elif val:
            result[base] += "\n" + "\n".join(val)
    return result


def extract_candidate_name(text: str) -> str:
    blocked = {"resume", "curriculum", "vitae", "summary", "profile", "experience", "education", "skills", "contact", "objective"}
    for line in [line.strip() for line in text.splitlines() if line.strip()][:8]:
        if any(marker in line.lower() for marker in ("@", "http", "linkedin", "github")):
            continue
        words = re.findall(r"[A-Za-z][A-Za-z'.-]*", line)
        if (2 <= len(words) <= 5 and len(line) <= 60
                and not blocked.intersection(word.lower() for word in words)
                and sum(word[0].isupper() for word in words) >= len(words) - 1):
            return " ".join(words)
    return "Unknown Candidate"
